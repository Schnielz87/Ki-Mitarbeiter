"""Die Bedienungsanleitung darf nichts versprechen, was es nicht gibt.

Die Anleitung wird aus `tools/anleitung_erzeugen.py` erzeugt. Sie nennt
Befehle der Konsolenfassung und Ordner auf dem Datentraeger. Beides kann
sich im Programm aendern, ohne dass jemand daran denkt, die Anleitung
nachzuziehen - dann steht dort eine Anweisung, die ins Leere laeuft.

Dieser Test vergleicht deshalb den Text der Anleitung mit dem Programm.
"""

from __future__ import annotations

import re
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
QUELLE = (ROOT / "tools" / "anleitung_erzeugen.py").read_text(encoding="utf-8")

#: Der Programmname taucht auch im Fliesstext auf ("... KONSOLE.exe
#: verwenden - dieselbe Anwendung"). Solche Stellen sind keine Befehle; sie
#: erkennt man daran, dass ein Gedankenstrich folgt.
_BEFEHL = re.compile(r"PORTABLE_BUCHHALTER_KONSOLE\.exe\s+([a-z]+)(?![a-z])(?!\s+-\s)")


def test_alle_genannten_befehle_gibt_es_wirklich():
    from ui.cli import build_parser

    parser = build_parser()
    unterbefehle = set()
    for aktion in parser._subparsers._group_actions:      # noqa: SLF001 - Testzugriff
        unterbefehle.update(aktion.choices)

    genannt = set(_BEFEHL.findall(QUELLE))
    erfunden = sorted(genannt - unterbefehle)
    assert not erfunden, (
        "Die Anleitung nennt Befehle, die es nicht gibt: " + ", ".join(erfunden)
    )


def test_die_wichtigen_neuen_befehle_stehen_in_der_anleitung():
    """Was gebaut wurde, muss auch erklaert sein - sonst findet es niemand."""
    genannt = set(_BEFEHL.findall(QUELLE))
    for befehl in ("datei", "plugin", "quellen", "modus"):
        assert befehl in genannt, f"Der Befehl '{befehl}' fehlt in der Anleitung."


def test_genannte_ordner_stimmen_mit_dem_layout_ueberein():
    from pkc.paths import LAYOUT

    for name in ("workspace\\\\artefakte", "workspace/artefakte"):
        if name in QUELLE:
            break
    else:
        raise AssertionError("Die Anleitung sagt nicht, wo erzeugte Dateien liegen.")
    assert LAYOUT["artefakte"] == "workspace/artefakte", \
        "Der Ablageort hat sich geaendert - die Anleitung muss nachgezogen werden."


def test_die_anleitung_liegt_vor():
    ziel = ROOT / "docs" / "BEDIENUNGSANLEITUNG.docx"
    assert ziel.is_file() and ziel.stat().st_size > 20_000, \
        "docs/BEDIENUNGSANLEITUNG.docx fehlt oder ist unvollstaendig."


def test_modellgroessen_stehen_nicht_fest_im_text():
    """Zahlen in der Anleitung muessen aus dem Katalog kommen, nicht aus dem Kopf.

    Vorher standen "0,4 GB", "4,7 GB" und "9 GB" als Text im Erzeuger. Der
    Bauablauf hat spaeter gemessen, dass es 0,49, 4,68 und 8,99 GB sind -
    die Anleitung war damit falsch, ohne dass es jemandem auffiel. Jetzt
    baut sie die Tabelle aus `config/model_catalog.json`.
    """
    assert "_katalog.laden(REPO" in QUELLE, (
        "Die Modelltabelle der Anleitung muss aus dem Katalog kommen.")
    for erfunden in ('"0,4 GB"', '"4,7 GB"', '"9 GB"', '"2 GB"'):
        assert erfunden not in QUELLE, (
            f"{erfunden} steht wieder fest im Text - dann kann es veralten.")


def test_anleitung_nennt_die_gemessenen_groessen():
    """Und die erzeugte Datei muss diese Zahlen dann auch wirklich zeigen."""
    from docx import Document

    import sys
    sys.path.insert(0, str(ROOT / "src"))
    from pkc.llm import katalog as katalogmodul

    quellen = {q.profil: q for q in katalogmodul.laden(ROOT / "config")}
    assert quellen, "ohne Katalog sagt dieser Test nichts aus"

    tabelle = None
    for kandidat in Document(str(ROOT / "docs" / "BEDIENUNGSANLEITUNG.docx")).tables:
        if [z.text for z in kandidat.rows[0].cells][:2] == ["Auswahl", "Groesse"]:
            tabelle = kandidat
            break
    assert tabelle is not None, "Die Modelltabelle fehlt in der Anleitung."

    gesehen = {}
    for zeile in tabelle.rows[1:]:
        felder = [z.text for z in zeile.cells]
        gesehen[felder[0]] = (felder[1], felder[2])
    assert set(gesehen) == set(quellen), "Die Tabelle deckt sich nicht mit dem Katalog."

    for profil, quelle in quellen.items():
        groesse, ram = gesehen[profil]
        assert f"{quelle.groesse_gb:.2f}".replace(".", ",") in groesse, profil
        assert f"{quelle.min_ram_gb} GB RAM" == ram, profil
        if quelle.geteilt:
            assert f"{len(quelle.teile)} Teildateien" in groesse, (
                f"{profil}: die Teildateien muessen angekuendigt werden")


def test_anleitung_erklaert_die_registerkarte_sprachmodell():
    """Eine neue Registerkarte ohne Anleitung ist eine halbe Lieferung.

    Anlass ist die Rueckfrage "hast du das auch in der Betriebsanleitung
    ergaenzt?". Der Kasten mit den drei Klicks stand da bereits - der Rest
    des Kapitels sprach aber weiter von Konsolenschaltern.
    """
    from docx import Document

    text = "\n".join(p.text for p in
                     Document(str(ROOT / "docs" / "BEDIENUNGSANLEITUNG.docx")).paragraphs)
    tabellen = "\n".join(z.text for t in
                         Document(str(ROOT / "docs" / "BEDIENUNGSANLEITUNG.docx")).tables
                         for r in t.rows for z in r.cells)
    alles = text + "\n" + tabellen

    # Die Kernaussage auf die Rueckfrage "muss ich das immer laden?"
    assert "EINMAL geladen" in alles
    assert "Beim naechsten Start wird nichts nachgeladen." in alles
    # Der Weg fuer den zweiten Datentraeger
    assert "Vorhandene Modelldatei uebernehmen" in alles
    assert "KOPIERT, nicht nur verknuepft" in alles
    # Und die Registerkarte selbst, Feld fuer Feld
    for stueck in ("Lage auf diesem Rechner", "Modell ausprobieren",
                   "Lage neu pruefen", "Sprachmodell einrichten"):
        assert stueck in alles, f"die Anleitung erklaert '{stueck}' nicht"


def test_anleitung_nennt_nur_schaltflaechen_die_es_gibt():
    """Sonst sucht jemand einen Knopf, den niemand gebaut hat."""
    oberflaeche = (ROOT / "src" / "ui" / "tk_app.py").read_text(encoding="utf-8")
    for schaltflaeche in ("Sprachmodell einrichten", "Vorhandene Modelldatei uebernehmen",
                          "Lage neu pruefen", "Modell ausprobieren"):
        assert f'text="{schaltflaeche}"' in oberflaeche, (
            f"Die Anleitung nennt '{schaltflaeche}' - im Fenster gibt es das nicht.")


def test_anleitung_beantwortet_die_frage_nach_der_geschwindigkeit():
    """"Warum dauert das so lange, und geht es wie bei ChatGPT?"

    Die Frage kam aus dem Betrieb. Die Anleitung muss beides beantworten:
    warum es langsamer ist als ein Rechenzentrum (und dass das der Preis
    fuer den Offlinebetrieb ist) und was auf diesem Rechner hilft.
    """
    from docx import Document

    datei = str(ROOT / "docs" / "BEDIENUNGSANLEITUNG.docx")
    dokument = Document(datei)
    alles = "\n".join(p.text for p in dokument.paragraphs) + "\n" + "\n".join(
        z.text for t in dokument.tables for r in t.rows for z in r.cells)

    assert "ChatGPT" in alles, "die Frage wird gestellt - sie gehoert beantwortet"
    assert "Rechenzentren" in alles and "Buerorechner" in alles
    # Der Hauptgrund fuer die 0,3 Token je Sekunde
    assert "passt nicht in den Arbeitsspeicher" in alles
    assert "ZU GROSS FUER DIESEN RECHNER" in alles
    # Und der Weg zu echter Geschwindigkeit
    assert "Grafikschichten" in alles
    assert "runtime\\llama" in alles


def test_anleitung_verspricht_keine_zahlen_als_zusage():
    """Groessenordnungen ja - Zusagen nein. Gemessen wird auf dem Rechner."""
    from docx import Document

    alles = "\n".join(p.text for p in
                      Document(str(ROOT / "docs" / "BEDIENUNGSANLEITUNG.docx")).paragraphs)
    assert "keine Zusage" in alles
    assert "misst das auf Ihrem Rechner selbst" in alles
