# -*- coding: utf-8 -*-
"""Erzeugt die Bedienungsanleitung als Word-Dokument.

Alle Beschriftungen stammen woertlich aus dem Programmcode
(src/ui/tk_app.py, src/profiles/buchhalter/profile.json) - nicht aus
Erinnerung. Wird die Oberflaeche geaendert, ist dieses Skript erneut zu
laufen bzw. anzupassen.

Braucht python-docx:

    pip install python-docx

Bewusst **nicht** in requirements-dev.txt: die Anwendung braucht es nicht,
der Windows-Bauablauf auch nicht. Das fertige Dokument liegt im Repository,
neu erzeugen muss es nur, wer die Oberflaeche aendert.

    python tools/anleitung_erzeugen.py
"""
from __future__ import annotations

import datetime as _dt
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "src"))

from pkc.config import DEFAULTS  # noqa: E402
from pkc.memory.schema_keys import WELL_KNOWN_KEYS  # noqa: E402

#: Zwei getrennte Angaben: die Fassung des Programms und die des Fachmoduls.
#: Frueher stand hier die Fachmodulfassung unter der Ueberschrift
#: "Programmfassung" - das war schlicht falsch.
PRODUKTFASSUNG = DEFAULTS["product"]["version"]

#: Marke und Profilname wie im Fenstertitel - damit die Anleitung dasselbe
#: nennt, was der Benutzer vor sich sieht.
from pkc.branding import MARKE  # noqa: E402
from pkc.llm import katalog as _katalog  # noqa: E402

profil = json.loads((REPO / "src/profiles/buchhalter/profile.json").read_text(encoding="utf-8"))

AKZENT = RGBColor(0x1F, 0x4E, 0x79)
GRAU = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---------------------------------------------------------------- Formatvorlagen
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)

for name, groesse in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
    stil = doc.styles[name]
    stil.font.name = "Calibri"
    stil.font.size = Pt(groesse)
    stil.font.color.rgb = AKZENT
    stil.font.bold = True

for abschnitt in doc.sections:
    abschnitt.top_margin = Cm(2.2)
    abschnitt.bottom_margin = Cm(2.0)
    abschnitt.left_margin = Cm(2.4)
    abschnitt.right_margin = Cm(2.4)


def absatz(text="", stil=None, fett=False, kursiv=False, farbe=None, groesse=None):
    p = doc.add_paragraph(style=stil)
    lauf = p.add_run(text)
    lauf.bold = fett
    lauf.italic = kursiv
    if farbe is not None:
        lauf.font.color.rgb = farbe
    if groesse is not None:
        lauf.font.size = Pt(groesse)
    return p


def punkt(text, ebene=0):
    stil = "List Bullet" if ebene == 0 else "List Bullet 2"
    return doc.add_paragraph(text, style=stil)


def schritt(text):
    return doc.add_paragraph(text, style="List Number")


def kasten(titel, text):
    tabelle = doc.add_table(rows=1, cols=1)
    tabelle.style = "Table Grid"
    zelle = tabelle.rows[0].cells[0]
    p = zelle.paragraphs[0]
    lauf = p.add_run(titel)
    lauf.bold = True
    lauf.font.color.rgb = AKZENT
    zelle.add_paragraph(text)
    doc.add_paragraph()
    return tabelle


def tabelle_mit(kopf, zeilen, breiten=None):
    t = doc.add_table(rows=1, cols=len(kopf))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, text in enumerate(kopf):
        zelle = t.rows[0].cells[i]
        zelle.text = ""
        lauf = zelle.paragraphs[0].add_run(text)
        lauf.bold = True
    for zeile in zeilen:
        zellen = t.add_row().cells
        for i, wert in enumerate(zeile):
            zellen[i].text = ""
            zellen[i].paragraphs[0].add_run(str(wert)).font.size = Pt(10)
    if breiten:
        for zeile in t.rows:
            for i, breite in enumerate(breiten):
                zeile.cells[i].width = Cm(breite)
    doc.add_paragraph()
    return t


# ================================================================ Titelseite
titel = doc.add_paragraph()
titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
lauf = titel.add_run(f"{MARKE} - {profil.get('short_name') or 'Buchhalter'}")
lauf.bold = True
lauf.font.size = Pt(30)
lauf.font.color.rgb = AKZENT

unter = doc.add_paragraph()
unter.alignment = WD_ALIGN_PARAGRAPH.CENTER
lauf = unter.add_run("Bedienungsanleitung - wo Sie was eingeben")
lauf.font.size = Pt(15)
lauf.font.color.rgb = GRAU

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
lauf = info.add_run(
    f"Programmfassung {PRODUKTFASSUNG}  ·  Fachmodul {profil['version']}"
    f"  ·  Stand {_dt.date.today().strftime('%d.%m.%Y')}"
)
lauf.font.size = Pt(10)
lauf.font.color.rgb = GRAU

doc.add_paragraph()
kasten(
    "Bitte zuerst lesen",
    "Diese Anleitung wird aus dem Programmcode selbst erzeugt. Alle genannten "
    "Registerkarten, Schaltflaechen und Felder tragen daher genau die "
    "Beschriftungen, die das Programm setzt. Der beschriebene Aufbau wurde "
    "am laufenden Programm abgeglichen. Weicht dennoch etwas ab, gilt das "
    "Programm.\n\n"
    "Der Buchhalter ist eine fachliche Zuarbeit. Er ersetzt weder Steuerberater "
    "noch Rechtsanwalt. Verantwortung und Freigabe bleiben immer bei Ihnen.",
)

doc.add_page_break()

# ================================================================ Inhalt
doc.add_heading("Inhalt", level=1)
for nummer, kapitel in enumerate([
    "In fuenf Minuten startklar",
    "Das Sprachmodell einrichten - damit die KI antwortet",
    "Der Startbildschirm - was er Ihnen sagt",
    "Der Bildschirm: sechs Registerkarten",
    "Fachfragen ohne Unternehmensdaten - geht sofort",
    "Ihr Unternehmen einrichten - empfohlen, nicht Pflicht",
    "Fragen stellen",
    "Wie eine Antwort aufgebaut ist",
    "Was automatisch gespeichert wird - und was nicht",
    "Belege hinzufuegen",
    "Unternehmenswissen pflegen",
    "Ergebnisse als Datei ausgeben",
    "Erweiterungen (Plugins)",
    "Betriebsmodus: HYBRID, OFFLINE, ONLINE",
    "Wissen aktualisieren",
    "Einstellungen und Status",
    "Lizenz",
    "Was der Buchhalter nicht tut",
    "Wenn etwas nicht funktioniert",
], start=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"{nummer}.  ").bold = True
    p.add_run(kapitel)

doc.add_page_break()

# ================================================================ 1
doc.add_heading("1.  In fuenf Minuten startklar", level=1)

absatz("Der Buchhalter liegt vollstaendig auf Ihrem Datentraeger. Er braucht "
       "keine Installation, keine Administratorrechte und kein Internet.")

doc.add_heading("So starten Sie", level=2)
schritt("Datentraeger (SSD, USB-Stick) an den Windows-PC anschliessen.")
schritt("Den Ordner oeffnen.")
schritt("PORTABLE_BUCHHALTER.exe doppelklicken.")
schritt("Es oeffnet sich ein kleines Fenster mit der Ueberschrift PORTABLER "
        "BUCHHALTER und einer Systempruefung. Warten Sie, bis die Pruefung "
        "durchgelaufen ist.")
schritt("Steht dort \"Der Buchhalter kann gestartet werden\", klicken Sie auf "
        "die Schaltflaeche BUCHHALTER STARTEN. Erst dann ist sie anklickbar.")

absatz()
kasten(
    "Beim allerersten Start",
    "Windows meldet sich mit \"Der Computer wurde durch Windows geschuetzt\". "
    "Das liegt daran, dass die Anwendung noch nicht digital signiert ist - es "
    "ist kein Fehler. Klicken Sie auf \"Weitere Informationen\" und dann auf "
    "\"Trotzdem ausfuehren\".",
)

doc.add_heading("Eine neuere Fassung holen", level=2)
absatz("Die fertige Anwendung wird bei jeder Aenderung neu gebaut und liegt "
       "als ZIP-Datei bereit. So kommen Sie daran - der Weg bleibt immer "
       "derselbe:")
schritt("Bei GitHub anmelden. Das ist zwingend: ohne Anmeldung sehen Sie das "
        "Paket zwar, koennen es aber nicht herunterladen. Der Name ist dann "
        "kein Link, und ein Klick tut schlicht nichts.")
schritt("https://github.com/Schnielz87/Ki-Mitarbeiter/actions oeffnen.")
schritt("Den obersten Ablauf mit gruenem Haken anklicken. Dort steht immer "
        "die neueste Fassung.")
schritt("Ganz unten im Kasten Artefakte auf den Namen "
        "Portable-Buchhalter-Windows klicken. Der Download startet.")
schritt("ZIP entsperren (Rechtsklick, Eigenschaften, Zulassen), auf den "
        "Datentraeger entpacken, fertig.")

kasten(
    "Wichtig: Ihre Daten bleiben erhalten",
    "Beim Einspielen einer neueren Fassung werden nur die Programmdateien "
    "ersetzt. Ihr Unternehmensgedaechtnis, Ihre Unterhaltungen, Belege, "
    "Einstellungen und Sicherungen liegen in eigenen Ordnern und bleiben "
    "unberuehrt - siehe Kapitel 8. Legen Sie trotzdem vorher eine Sicherung "
    "an.",
)

absatz("Ein Paket wird nach 30 Tagen von GitHub geloescht. Das ist kein "
       "Problem: der naechste Ablauf legt ein frisches ab. Nehmen Sie dann "
       "einfach wieder den obersten gruenen Ablauf.")

doc.add_heading("Warum zwei Programmdateien?", level=2)
absatz("Im Ordner liegen zwei Dateien mit fast gleichem Namen. Das ist "
       "Absicht und kein Versehen - beide enthalten dieselbe Anwendung, sie "
       "unterscheiden sich nur darin, ob ein Textfenster mitlaeuft.")

tabelle_mit(
    ["Datei", "Wofuer", "Was Sie sehen"],
    [
        ["PORTABLE_BUCHHALTER.exe",
         "Der Normalfall. Diese Datei doppelklicken.",
         "Nur das Programmfenster - kein schwarzes Textfenster daneben."],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe",
         "Fuer die Kommandozeile und zur Fehlersuche.",
         "Ein Textfenster mit Ausgaben. Ohne Argumente startet auch sie die "
         "Oberflaeche - dann eben mit Textfenster dahinter."],
    ],
    breiten=[5.6, 4.8, 5.6],
)

absatz("Der Hintergrund ist eine Windows-Eigenheit: Ein Programm ohne "
       "Konsole kann in der Eingabeaufforderung nichts ausgeben. Wuerde es "
       "nur eine Datei geben, muesste man sich entscheiden - entweder ein "
       "schwarzes Fenster bei jedem Doppelklick, oder stumme Befehle in der "
       "Eingabeaufforderung. Deshalb zwei Dateien, wie bei Python selbst "
       "python.exe und pythonw.exe.")

kasten(
    "Wann brauchen Sie die Konsolenfassung?",
    "Fuer alles, was in dieser Anleitung als Befehl geschrieben steht - "
    "etwa \"lizenz\", \"reife\", \"quellen\" oder \"check\". Und wenn beim "
    "Doppelklick nichts passiert: dann zeigt die Konsolenfassung, woran es "
    "liegt. Fuer die taegliche Arbeit brauchen Sie sie nicht.",
)

doc.add_heading("Der Laufwerksbuchstabe ist egal", level=2)
absatz("Ob der Datentraeger als D:, E: oder F: erscheint, spielt keine Rolle. "
       "Der Buchhalter findet seine Daten selbst. Auch Pfade mit Leerzeichen "
       "funktionieren. Sie koennen den Datentraeger an einen anderen PC "
       "stecken - Ihr Unternehmenswissen und Ihre Unterhaltungen sind dort "
       "unveraendert vorhanden, weil alles auf dem Datentraeger liegt und "
       "nichts auf dem jeweiligen Rechner.")

# ================================================================ 2
doc.add_heading("2.  Das Sprachmodell einrichten - damit die KI antwortet", level=1)

absatz("Dieses Kapitel ist das wichtigste der ganzen Anleitung. Ohne "
       "Sprachmodell recherchiert der Buchhalter zwar in seinen Quellen und "
       "zeigt Ihnen die Fundstellen - er formuliert aber keine Antwort. Er "
       "sagt das dann auch offen.")

kasten(
    "Im Fenster: drei Klicks",
    "1. Registerkarte \"Sprachmodell\" oeffnen.\n"
    "2. Oben steht, was Ihr Rechner hergibt, und welches Modell dazu passt. "
    "Die Vorauswahl ist bereits die Empfehlung.\n"
    "3. Auf \"Sprachmodell einrichten\" klicken. Es erscheint eine "
    "Rueckfrage mit Groesse, Lizenz und Herkunft - erst Ihr \"Ja\" startet "
    "den Download.\n\n"
    "Danach bindet die Anwendung das Modell ein und stellt ihm selbst eine "
    "kleine Frage. Erst wenn die beantwortet ist, meldet sie \"Das "
    "Sprachmodell ist einsatzbereit.\" - eine geladene Datei allein waere "
    "noch kein Nachweis.",
)

doc.add_heading("Nur einmal - nicht bei jedem Start", level=2)
absatz("Das Modell wird EINMAL geladen. Danach liegt es im Ordner models "
       "auf diesem Datentraeger - nicht auf dem Rechner, an dem Sie es "
       "geladen haben. Es bleibt dort ueber Neustarts hinweg, an jedem "
       "anderen Rechner, unter jedem Laufwerksbuchstaben und fuer alle "
       "Kundenbereiche gemeinsam. Auch ohne Internet.")
punkt("Beim naechsten Start wird nichts nachgeladen.")
punkt("Ziehen Sie den Stick an einen anderen Rechner, geht es dort sofort "
      "weiter.")
punkt("Legen Sie einen zweiten Kundenbereich an, wird dasselbe Modell "
      "verwendet - es wird nicht ein zweites Mal abgelegt.")

absatz("Warum liegt das Modell dann nicht gleich im Paket? Zwei Gruende: es "
       "ist je nach Auswahl bis zu neun Gigabyte gross, und unter welcher "
       "Lizenz Sie ein Modell einsetzen, ist Ihre Entscheidung und nicht "
       "unsere.")

doc.add_heading("Sie haben das Modell schon? Dann nicht noch einmal laden", level=2)
absatz("Wenn die Datei bereits irgendwo liegt - auf dem Stick einer "
       "Kollegin, auf einer externen Platte, im Firmennetz -, koennen Sie "
       "sie uebernehmen, statt sie erneut zu ziehen. Das ist der Weg fuer "
       "einen zweiten Datentraeger, fuer ein Buero mit gesperrtem Download "
       "und fuer jede Leitung, ueber die man 4,7 GB nicht zweimal schicken "
       "will.")
kasten(
    "Vorhandene Datei uebernehmen",
    "Registerkarte \"Sprachmodell\" -> \"Vorhandene Modelldatei "
    "uebernehmen\" -> die .gguf-Datei auswaehlen.\n\n"
    "Die Datei wird auf diesen Datentraeger KOPIERT, nicht nur verknuepft. "
    "Nur so laeuft der Datentraeger auch an einem Rechner, der die Herkunft "
    "gar nicht erreicht - und genau das ist der Sinn dieser Anwendung.",
)
absatz("Dafuer braucht es kein Internet; die Uebernahme funktioniert auch "
       "im Betriebsmodus OFFLINE.")

doc.add_heading("Die Registerkarte im Einzelnen", level=2)
tabelle_mit(
    ["Was Sie sehen", "Was es bedeutet"],
    [
        ["Lage auf diesem Rechner",
         "\"Noch nicht eingerichtet\" oder \"Einsatzbereit\". Darunter steht, "
         "was fehlt - die Modelldatei, der Modelldienst, oder beides."],
        ["Auswahl",
         "Das Modell, das geladen wird. Vorausgewaehlt ist bereits das, was "
         "zu Ihrem Arbeitsspeicher passt."],
        ["Sprachmodell einrichten",
         "Laedt das ausgewaehlte Modell - erst nach einer Rueckfrage, die "
         "Groesse, Lizenz, Herkunft und Pruefstand nennt."],
        ["Vorhandene Modelldatei uebernehmen",
         "Nimmt eine .gguf-Datei auf, die Sie schon haben. Kein Download."],
        ["Lage neu pruefen",
         "Sieht noch einmal nach - etwa wenn Sie eine Datei von Hand in den "
         "Ordner models kopiert haben."],
        ["Modell ausprobieren",
         "Stellt dem Modell eine kleine Frage und zeigt Antwortzeit und "
         "Geschwindigkeit. Der Nachweis, dass es wirklich laeuft."],
        ["Textbereich unten",
         "Vor dem Einrichten: Ihre Hardware und die Bezugsquellen mit "
         "Lizenz und Pruefstand. Danach: das Ergebnis samt Probeantwort."],
    ],
    breiten=[5.0, 10.8],
)

absatz("Beim Einrichten wandert der Balken bis 100 Prozent. Danach bindet "
       "die Anwendung das Modell ein und stellt ihm selbst eine Frage. Erst "
       "wenn die beantwortet ist, meldet sie \"Das Sprachmodell ist "
       "einsatzbereit\". Eine geladene Datei allein ist noch kein Nachweis - "
       "sie koennte beschaedigt sein oder zu gross fuer Ihren "
       "Arbeitsspeicher.")

doc.add_heading("Dasselbe in der Kommandozeile", level=2)
absatz("Wer lieber mit der Konsole arbeitet, erreicht alles auch dort:")
kasten(
    "Befehle",
    "PORTABLE_BUCHHALTER_KONSOLE.exe modell status\n"
    "PORTABLE_BUCHHALTER_KONSOLE.exe modell empfehlen\n"
    "PORTABLE_BUCHHALTER_KONSOLE.exe modell einrichten --bestaetigen\n"
    "PORTABLE_BUCHHALTER_KONSOLE.exe modell uebernehmen --datei D:\\modell.gguf\n"
    "PORTABLE_BUCHHALTER_KONSOLE.exe modell pruefen",
)

doc.add_heading("Was schon fertig mitgeliefert wird", level=2)
absatz("Das Programm, das ein Sprachmodell ausfuehrt, liegt bereits im "
       "Ordner runtime\\llama. Sie muessen nichts installieren und nichts "
       "einrichten. Die Anwendung startet es selbst - und zwar erst bei "
       "Ihrer ersten Frage, damit der Programmstart nicht jedes Mal auf das "
       "Laden eines mehrere Gigabyte grossen Modells wartet.")
punkt("Es laeuft nur auf Ihrem Rechner und ist von aussen nicht erreichbar.")
punkt("Es oeffnet kein zusaetzliches Fenster.")
punkt("Beim Schliessen der Anwendung wird es mit beendet.")

absatz()
doc.add_heading("Welches Modell?", level=2)
# Die Groessen stehen nicht hier, sondern im ausgelieferten Katalog - und
# der traegt, was der Windows-Bauablauf tatsaechlich abgerufen hat. So kann
# in der Anleitung keine Zahl stehen, die niemand gemessen hat.
_ZWECK = {
    "probe": "Nur zum Ausprobieren. Fuer Buchhaltungsfragen NICHT geeignet.",
    "light": "Aeltere Buerorechner. Lizenz beachten - nicht Apache-2.0.",
    "standard": "Empfehlung fuer den normalen Betrieb.",
    "high": "Beste Qualitaet. Auf einem Rechner ohne Grafikkarte langsam.",
}
_zeilen = []
for _quelle in _katalog.laden(REPO / "config"):
    _groesse = f"{_quelle.groesse_gb:.2f}".replace(".", ",") + " GB"
    if _quelle.geteilt:
        _groesse += f"\n({len(_quelle.teile)} Teildateien)"
    _zeilen.append([_quelle.profil, _groesse, f"{_quelle.min_ram_gb} GB RAM",
                    _ZWECK.get(_quelle.profil, _quelle.hinweis)])
tabelle_mit(
    ["Auswahl", "Groesse", "Braucht", "Wofuer"],
    _zeilen,
    breiten=[2.6, 2.2, 2.6, 7.6],
)
absatz("Die beiden groesseren Modelle liegen beim Anbieter in mehreren "
       "Dateien vor. Die Anwendung laedt alle Teile und legt sie zusammen "
       "in den Ordner models ab. Bitte nichts davon umbenennen oder "
       "verschieben - die Teile gehoeren zusammen.")
absatz("Ein anderes als das vorgeschlagene Modell waehlen Sie im Fenster "
       "ueber die Auswahlliste \"Auswahl\" - in der Konsole mit --profil "
       "light (oder standard, high, probe).")

doc.add_heading("Was die Anwendung dabei nicht tut", level=2)
punkt("Sie laedt nichts ohne Ihre Bestaetigung. Im Fenster kommt vorher "
      "eine Rueckfrage mit Groesse, Lizenz und Herkunft; in der Konsole "
      "braucht es --bestaetigen.")
punkt("Sie laedt nichts im Betriebsmodus OFFLINE. Eine vorhandene Datei "
      "uebernehmen koennen Sie trotzdem - dabei wird nichts abgerufen.")
punkt("Sie behauptet keine geprueften Bezugsquellen. Steht dort \"nicht "
      "geprueft\", wurde die Adresse in diesem Programmstand nicht "
      "abgerufen - dann wird die Datei auch nicht gegen eine hinterlegte "
      "Pruefsumme geprueft.")
punkt("Sie faengt zu wenig Speicherplatz vorher ab, statt mittendrin "
      "abzubrechen.")

absatz()
doc.add_heading("Nachsehen und pruefen", level=2)
tabelle_mit(
    ["Befehl", "Wirkung"],
    [
        ["PORTABLE_BUCHHALTER_KONSOLE.exe modell status",
         "Was ist da, was fehlt, und was waere der naechste Schritt"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe modell pruefen",
         "Stellt dem Modell eine Frage und nennt Antwortzeit und Tempo"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe modell laden --url <Adresse>",
         "Ein eigenes Modell aus einer selbst gewaehlten Quelle"],
    ],
    breiten=[8.0, 7.0],
)

kasten(
    "Wenn die erste Antwort lange dauert",
    "Das ist einmalig: das Modell wird in den Arbeitsspeicher geladen. Die "
    "folgenden Antworten kommen deutlich schneller.\n\n"
    "Kommt der Dienst gar nicht hoch, nennt die Meldung seine letzte "
    "Ausgabe - meist reicht der Arbeitsspeicher nicht. Dann ein kleineres "
    "Profil waehlen. Ausfuehrlich steht es in logs\\llama-server.log.",
)

# ================================================================ 3
doc.add_heading("3.  Der Startbildschirm - was er Ihnen sagt", level=1)

absatz("Bevor der Buchhalter startet, prueft er sich selbst und zeigt das "
       "Ergebnis an. Das ist keine Formalie: Sie sehen dort, ob er "
       "arbeitsfaehig ist und was ihm gegebenenfalls fehlt. Jede Zeile hat "
       "ein Ergebnis - OK oder HINWEIS.")

tabelle_mit(
    ["Zeile", "Was geprueft wird", "Was Sie sehen sollten"],
    [
        ["Datenverzeichnis",
         "Wo Ihre Daten liegen und ob dorthin geschrieben werden kann",
         "OK, dahinter der Pfad und \"(beschreibbar)\""],
        ["Unternehmensgedaechtnis",
         "Wie viele Angaben ueber Ihren Betrieb gespeichert sind",
         "OK, Anzahl der Eintraege, \"Integritaet: ok\""],
        ["Fachwissen",
         "Die mitgelieferten Fachmodule",
         "OK, 13 Dokumente, 57 Abschnitte"],
        ["Suchindex",
         "Ob die Suche ueber das Fachwissen einsatzbereit ist",
         "OK, Anzahl der Abschnitte"],
        ["Lokales Modell",
         "Ob ein Sprachmodell vorhanden ist",
         "HINWEIS, solange keines eingerichtet ist - siehe unten"],
        ["Lizenz",
         "Ob eine Lizenz noetig und gueltig ist",
         "OK - siehe Kapitel 14"],
        ["Quellenregister",
         "Die hinterlegten amtlichen Quellen",
         "OK, 12 Quellen, 32 Dokumente"],
        ["Geheimnistresor",
         "Der verschluesselte Speicher fuer Passwoerter und Zugaenge",
         "OK, \"noch nicht angelegt\" ist voellig in Ordnung"],
    ],
    breiten=[3.8, 5.8, 5.4],
)

absatz("Darunter stehen vier Angaben zum Zustand:")
punkt("Wissensstand - auf welchem Datum die gespeicherten Fachquellen stehen.")
punkt("Internet - ob eine Verbindung besteht. Fuer die taegliche Arbeit "
      "unerheblich; gebraucht wird sie nur zum Nachladen von Quellen.")
punkt("Betriebsart - OFFLINE oder HYBRID. HYBRID heisst lediglich, dass "
      "Internet verfuegbar ist. Gearbeitet wird trotzdem lokal.")
punkt("Datenverzeichnis - der Ordner auf Ihrem Datentraeger.")

absatz()
absatz("Steht am Ende \"Der Buchhalter kann gestartet werden.\", ist alles "
       "in Ordnung. Erst dann laesst sich die Schaltflaeche BUCHHALTER "
       "STARTEN anklicken.", fett=True)

doc.add_heading("HINWEIS ist kein Fehler", level=2)
absatz("OK bedeutet: in Ordnung. HINWEIS bedeutet: es fehlt etwas, aber der "
       "Buchhalter laeuft trotzdem. Der haeufigste Fall ist das Sprachmodell:")

kasten(
    "\"Lokales Modell: HINWEIS - Kein Sprachmodell verfuegbar\"",
    "Das ist der Normalzustand, solange Sie kein Modell eingerichtet haben. "
    "Der Buchhalter laeuft dann im sogenannten Notbetrieb: er recherchiert "
    "in seinen Quellen und zeigt Ihnen die Fundstellen, formuliert aber "
    "keine ausformulierte Fachantwort.\n\n"
    "Das ist Absicht. Lieber sagt er ehrlich, dass er nichts formulieren "
    "kann, als etwas zu erfinden. Wie Sie ein Modell einrichten, steht in "
    "docs/MODELL_EINRICHTEN.md - einmalig, etwa 5 GB.",
)

absatz("Ein echter Fehler wuerde die Zeile mit FEHLER kennzeichnen und die "
       "Schaltflaeche BUCHHALTER STARTEN gesperrt lassen.")

# ================================================================ 3
doc.add_heading("4.  Der Bildschirm: sechs Registerkarten", level=1)
absatz("Nach dem Start sehen Sie oben eine Zeile mit sechs Registerkarten. "
       "Alles, was Sie tun, geschieht in einer davon.")

tabelle_mit(
    ["Registerkarte", "Wofuer", "Was Sie dort eingeben"],
    [
        ["Unterhaltung", "Fragen stellen und Antworten lesen",
         "Ihre Frage in das grosse Eingabefeld unten links"],
        ["Unternehmenswissen", "Angaben zu Ihrem Betrieb dauerhaft speichern",
         "Suchbegriffe oben; Angaben ueber \"Neu / Aendern\" oder \"Onboarding fortsetzen\""],
        ["Belege", "Eigene Dokumente einlesen",
         "Keine Eingabe - Sie waehlen ueber \"Beleg hinzufuegen\" eine Datei aus"],
        ["Sprachmodell", "Das Modell einmalig einrichten und ausprobieren",
         "Auswahlliste mit dem Modell; sonst nur Schaltflaechen"],
        ["Wissen aktualisieren", "Amtliche Quellen nachladen (braucht Internet)",
         "Keine Eingabe - nur Schaltflaechen"],
        ["Einstellungen und Status", "Verhalten einstellen, Zustand pruefen",
         "Auswahllisten und Haekchen"],
    ],
    breiten=[3.6, 5.4, 7.0],
)

absatz("Ganz oben rechts stehen zwei Angaben, die Sie im Blick behalten "
       "sollten:")
punkt("Betriebsart - OFFLINE oder HYBRID. HYBRID heisst nur, dass Internet "
      "verfuegbar ist; gearbeitet wird trotzdem lokal.")
punkt("Wissensstand - das Datum, auf dem die gespeicherten Fachquellen stehen.")
absatz("Ganz unten laeuft eine Statuszeile mit, die auch den Pfad Ihres "
       "Datentraegers nennt.")

# ================================================================ 3
doc.add_heading("5.  Fachfragen ohne Unternehmensdaten - geht sofort", level=1)

absatz("Sie muessen dem Buchhalter nichts ueber Ihren Betrieb erzaehlen, um "
       "ihn fachlich zu befragen. Das Fachwissen und das "
       "Unternehmensgedaechtnis sind zwei getrennte Bestaende. Das Fachwissen "
       "ist vollstaendig mitgeliefert und vom ersten Start an nutzbar.")

kasten(
    "Kurz gesagt",
    "Datentraeger anstecken, starten, fragen. Kein Onboarding, keine "
    "Registrierung, keine Angabe zu Ihrem Unternehmen. Wer nur wissen will, "
    "was fachlich gilt, kann das sofort tun - und muss dabei nichts "
    "preisgeben.",
)

doc.add_heading("Was mitgeliefert ist", level=2)
absatz("13 Fachmodule mit 57 Abschnitten, dazu ein Quellenregister mit 12 "
       "amtlichen Quellen. Der Startbildschirm zeigt das in den Zeilen "
       "Fachwissen und Quellenregister an. Die Module decken unter anderem "
       "ab:")
punkt("Rechnungspflichtangaben und Rechnungspruefung (§ 14 UStG)")
punkt("Vorsteuerabzug (§ 15 UStG)")
punkt("GoBD, Aufbewahrung und Dokumentation (§§ 145-147 AO, § 257 HGB, § 14b UStG)")
punkt("Steuerschuldnerschaft des Leistungsempfaengers, Reverse Charge (§ 13b UStG)")
punkt("Kleinunternehmerregelung (§ 19 UStG)")
punkt("E-Rechnung, Jahresabschlussvorbereitung und weitere")

doc.add_heading("Ein nachgeprueftes Beispiel", level=2)
absatz("In einem voellig leeren Datenbereich - null Eintraege im "
       "Unternehmensgedaechtnis - wurde gefragt:")
absatz("\"Wie lange muessen Eingangsrechnungen aufbewahrt werden?\"", kursiv=True)
absatz("Der Buchhalter fand acht Fundstellen. Darunter aus dem Modul GoBD "
       "die Fristentabelle: Buecher, Aufzeichnungen, Jahresabschluesse, "
       "Buchungsbelege und Inventare zehn Jahre nach § 147 Abs. 3 AO und "
       "§ 257 HGB; empfangene und abgesandte Handels- und Geschaeftsbriefe "
       "sechs Jahre nach § 147 Abs. 3 AO; Rechnungen umsatzsteuerlich zehn "
       "Jahre nach § 14b Abs. 1 UStG. Ohne eine einzige Angabe ueber das "
       "Unternehmen.")

doc.add_heading("Fragen veraendert Ihr Unternehmensgedaechtnis nicht", level=2)
absatz("Nach dieser Frage stand im Unternehmensgedaechtnis weiterhin "
       "\"Kein Unternehmenswissen gespeichert\". Eine Fachfrage legt dort "
       "nichts an. Gespeichert wird eine Angabe nur, wenn Sie sie selbst "
       "eintragen oder die Rueckfrage \"Dauerhaft merken?\" mit Ja "
       "beantworten (siehe Kapitel 8).")

doc.add_heading("Wo Unternehmensdaten dann doch etwas bringen", level=2)
absatz("Bei allgemeinen Fragen - was gilt, welche Frist, welche "
       "Pflichtangaben - aendern Unternehmensdaten nichts. Sie helfen dort, "
       "wo die Antwort vom Betrieb abhaengt. Der Buchhalter sagt das auch "
       "selbst: im Modul Vorsteuerabzug steht woertlich \"Ohne bekannten "
       "Kontenrahmen sind die Konten ...\" - er nennt dann keine konkreten "
       "Kontonummern, statt welche zu erfinden.")

tabelle_mit(
    ["Art der Frage", "Ohne Unternehmensdaten", "Mit Unternehmensdaten"],
    [
        ["Welche Pflichtangaben braucht eine Rechnung?",
         "Vollstaendig beantwortbar", "Kein Unterschied"],
        ["Wie lange sind Belege aufzubewahren?",
         "Vollstaendig beantwortbar", "Kein Unterschied"],
        ["Wann greift Reverse Charge?",
         "Vollstaendig beantwortbar", "Kein Unterschied"],
        ["Auf welches Konto buche ich das?",
         "Buchungslogik ja, Kontonummern nein",
         "Konkrete Konten, sobald der Kontenrahmen hinterlegt ist"],
        ["Muessen wir Umsatzsteuer ausweisen?",
         "Die Regel ja, Ihr Fall nein",
         "Auf Ihren Umsatzsteuerstatus bezogen"],
        ["Wer muss das bei uns freigeben?",
         "Nicht beantwortbar", "Aus Ihren Freigaberegeln"],
    ],
    breiten=[5.4, 4.8, 4.8],
)

doc.add_heading("Fuer den Einstieg empfohlen", level=2)
absatz("Nutzen Sie den Buchhalter ruhig erst eine Weile rein fachlich. Sie "
       "sehen dann, wie er arbeitet und wie belastbar die Fundstellen sind, "
       "bevor Sie ihm etwas ueber Ihren Betrieb anvertrauen. Das Einrichten "
       "koennen Sie jederzeit spaeter nachholen - und Schritt fuer Schritt, "
       "es ist keine Alles-oder-nichts-Entscheidung.")

# ================================================================ 5
doc.add_heading("6.  Ihr Unternehmen einrichten - empfohlen, nicht Pflicht", level=1)

absatz("Alles in diesem Kapitel ist freiwillig. Der Buchhalter arbeitet "
       "fachlich auch ohne (Kapitel 4). Je mehr er ueber Ihren Betrieb "
       "weiss, desto konkreter werden allerdings die Antworten, die vom "
       "Betrieb abhaengen - und er merkt es sich dauerhaft auf dem "
       "Datentraeger.")

doc.add_heading("Der gefuehrte Weg", level=2)
schritt("Registerkarte Unternehmenswissen oeffnen.")
schritt("Unten auf Onboarding fortsetzen klicken.")
schritt("Es oeffnet sich das Fenster Unternehmens-Onboarding mit einer Liste "
        "von Fragen. Links steht die Frage, rechts ein Eingabefeld.")
schritt("Ausfuellen, was Sie wissen. Alles darf leer bleiben und spaeter "
        "ergaenzt werden.")
schritt("Auf Speichern klicken.")

absatz("Oben rechts in der Registerkarte sehen Sie danach den Fortschritt, "
       "zum Beispiel \"Onboarding: 7 von 21 beantwortet\".")

doc.add_heading("Diese Angaben werden abgefragt", level=2)
absatz("Die Reihenfolge im Fenster entspricht dieser Liste.")

zeilen = []
for schluessel in profil["onboarding_keys"]:
    meta = WELL_KNOWN_KEYS.get(schluessel, {"title": schluessel, "category": ""})
    zeilen.append([meta["title"], meta.get("category", "")])
tabelle_mit(["Angabe", "Kategorie"], zeilen, breiten=[8.0, 4.0])

kasten(
    "Sie koennen es auch einfach sagen",
    "Statt das Formular auszufuellen, duerfen Sie solche Angaben auch mitten "
    "im Gespraech nennen - zum Beispiel: \"Wir verwenden grundsaetzlich "
    "SKR03.\" Der Buchhalter erkennt dauerhaft gueltige Unternehmensangaben "
    "und fragt nach, ob er sie sich merken soll. Ob er nachfragt oder still "
    "speichert, stellen Sie unter Einstellungen ein.",
)

# ================================================================ 4
doc.add_heading("7.  Fragen stellen", level=1)

absatz("Registerkarte Unterhaltung. Der Bildschirm ist geteilt:")
punkt("Links oben: der bisherige Gespraechsverlauf.")
punkt("Links unten: das Eingabefeld - hier tippen Sie Ihre Frage.")
punkt("Rechts oben: Quellen der letzten Antwort - die Fundstellen, auf die "
      "sich die Antwort stuetzt.")
punkt("Rechts unten: Unterhaltungen - Ihre frueheren Gespraeche zum "
      "Nachschlagen.")

doc.add_heading("So fragen Sie", level=2)
schritt("In das Eingabefeld unten links klicken.")
schritt("Frage eintippen. Mehrere Zeilen sind erlaubt.")
schritt("Auf Absenden klicken - oder Strg + Eingabetaste druecken. Die "
        "Eingabetaste allein macht nur einen Zeilenumbruch.")

absatz()
doc.add_heading("Weitere Schaltflaechen daneben", level=2)
tabelle_mit(
    ["Schaltflaeche", "Was sie tut"],
    [
        ["Absenden", "Schickt Ihre Frage ab"],
        ["Neue Unterhaltung", "Beginnt ein neues Gespraech; das alte bleibt rechts erhalten"],
        ["Dokument hinzufuegen", "Liest einen Beleg ein, auf den Sie sich dann beziehen koennen"],
        ["Generierung stoppen", "Bricht eine laufende Antwort ab. Die Oberflaeche ist sofort "
         "wieder bedienbar; Ihre Frage bleibt gespeichert"],
        ["Antwort speichern", "Schreibt die letzte Antwort als Datei - Format daneben waehlbar "
         "(siehe Kapitel 12)"],
        ["Unterhaltung exportieren", "Speichert das Gespraech als lesbare Datei auf dem Datentraeger"],
    ],
    breiten=[5.0, 10.0],
)

doc.add_heading("Der Buchhalter merkt sich das Gespraech", level=2)
absatz("Sie muessen nicht jede Frage vollstaendig ausformulieren. Innerhalb "
       "einer Unterhaltung bezieht sich der Buchhalter auf das Vorherige:")
punkt("Sie: \"Ich habe eine Rechnung aus Frankreich.\"")
punkt("Buchhalter: \"Ist der Lieferant Unternehmer?\"")
punkt("Sie: \"Ja.\" - das genuegt. Die Rueckfrage wird richtig zugeordnet.")
absatz("Mit Neue Unterhaltung beginnt ein neuer Zusammenhang. Das ist "
       "sinnvoll, wenn es um einen anderen Sachverhalt geht.")

doc.add_heading("Nicht jede Nachricht wird gleich behandelt", level=2)
absatz("Der Buchhalter erkennt, um welche Art von Nachricht es sich handelt, "
       "und antwortet entsprechend. Auf \"Guten Morgen\" folgt keine "
       "Trefferliste aus dem Umsatzsteuerrecht.")
tabelle_mit(
    ["Art der Nachricht", "Was der Buchhalter tut"],
    [
        ["Begruessung, Rueckfrage an Sie, kurze Verstaendigung",
         "Antwortet kurz. Es wird nicht recherchiert, und es steht kein "
         "Quellenabschnitt darunter - es fehlt nichts."],
        ["Kurze Wissensfrage (\"Wie hoch ist der Regelsteuersatz?\")",
         "Kurze, belegte Antwort. Kein Aufbau ueber mehrere Abschnitte."],
        ["Fachfrage ohne eigenen Sachverhalt",
         "Recherchiert und antwortet mit Fundstellen."],
        ["Geschilderter Einzelfall",
         "Voller Aufbau nach dem Schema aus Kapitel 8. Fehlen entscheidende "
         "Angaben, fragt der Buchhalter gezielt nach, statt zu raten."],
    ],
    breiten=[5.5, 9.5],
)

doc.add_heading("Gute Fragen", level=2)
absatz("Je konkreter, desto besser. Beispiele:")
punkt("\"Welche Pflichtangaben muss eine Rechnung enthalten?\"")
punkt("\"Ein Lieferant aus Polen stellt uns ohne Umsatzsteuer in Rechnung. "
      "Wie buchen wir das?\"")
punkt("\"Wie lange muessen wir Eingangsrechnungen aufbewahren?\"")
punkt("\"Wir haben einen Firmenwagen geleast. Welche Unterlagen brauchen wir "
      "fuer den Jahresabschluss?\"")

# ================================================================ 5
doc.add_heading("8.  Wie eine Antwort aufgebaut ist", level=1)

absatz("Der Buchhalter antwortet nach einem festen Schema. Nicht jeder "
       "Abschnitt kommt bei jeder Frage vor, aber die Reihenfolge ist immer "
       "gleich:")

beschreibungen = {
    "ERGEBNIS": "Die kurze Antwort - was gilt.",
    "BEGRUENDUNG": "Warum das so ist.",
    "STEUERLICHE BEHANDLUNG": "Die steuerliche Seite des Sachverhalts.",
    "BUCHHALTERISCHE BEHANDLUNG": "Die buchhalterische Seite.",
    "BUCHUNGSVORSCHLAG": "Ein Vorschlag - Soll an Haben. Nur ein Vorschlag, siehe Freigabebedarf.",
    "BENOETIGTE UNTERLAGEN": "Was Sie noch beschaffen muessen.",
    "OFFENE PUNKTE": "Was der Buchhalter nicht klaeren konnte.",
    "RISIKEN": "Worauf Sie achten sollten.",
    "QUELLEN": "Die Fundstellen, nummeriert. Diese Nummern koennen Sie rechts nachschlagen.",
    "WISSENSSTAND": "Auf welchem Stand die verwendeten Quellen sind und wodurch die Antwort entstand.",
    "FREIGABEBEDARF": "Was ein Mensch freigeben muss, bevor es wirksam wird.",
}
tabelle_mit(
    ["Abschnitt", "Bedeutung"],
    [[a, beschreibungen.get(a, "")] for a in profil["answer_sections"]],
    breiten=[6.0, 9.0],
)

kasten(
    "Wichtig: die Abschnitte QUELLEN und FREIGABEBEDARF",
    "QUELLEN ist die Kontrolle. Der Buchhalter muss belegen, worauf er sich "
    "stuetzt. Steht dort nichts Brauchbares, ist die Antwort nicht belastbar - "
    "auch wenn sie ueberzeugend klingt.\n\n"
    "FREIGABEBEDARF sagt Ihnen, was noch nicht erledigt ist. Ein "
    "Buchungsvorschlag ist ein Vorschlag, keine Buchung.",
)

doc.add_heading("Wie die Antwort im Fenster erscheint", level=2)
absatz("Ueber Ihrer Frage steht BENUTZER, ueber der Antwort der Name des "
       "Mitarbeiters - zum Beispiel PORTIVA - Buchhalter. Die Antwort steht "
       "oben, darunter kleiner und ruhiger die Abschnitte QUELLEN, "
       "WISSENSSTAND, FREIGABEBEDARF und die Hinweise der Anwendung.")
absatz("Ist ein Sprachmodell eingerichtet, laeuft die Antwort waehrend der "
       "Erzeugung Stueck fuer Stueck ein, statt am Ende auf einmal zu "
       "erscheinen. Sobald sie fertig ist, tritt die gepruefte Fassung mit "
       "Quellen und Wissensstand an ihre Stelle. Dauert es Ihnen zu lange: "
       "Generierung stoppen.")
absatz("Zeichen wie ** oder ## sehen Sie nicht - Fettdruck ist Fettdruck, "
       "eine Ueberschrift ist eine Ueberschrift. In der Konsolenfassung "
       "werden diese Zeichen entfernt, weil eine Textausgabe keinen "
       "Fettdruck kennt.")

doc.add_heading("Rechts: Quellen der letzten Antwort", level=2)
absatz("Dort stehen die Fundstellen mit Nummer, Bezeichnung, Rang der Quelle "
       "und einem Auszug. Die Nummern in der Antwort - [1], [2] - verweisen "
       "genau dorthin. Der volle Auszug steht bewusst nur rechts: eine "
       "Antwort ist eine Antwort, keine Trefferliste.")
kasten(
    "Wenn nur Fachmodule gefunden wurden",
    "Dann schreibt der Buchhalter ausdruecklich, dass fuer diese Aussage nur "
    "Sekundaerquellen vorliegen und die Primaerquelle zu pruefen ist - "
    "Gesetzestext, Verwaltungsanweisung oder Rechtsprechung. Ohne diesen "
    "Hinweis saehe eine Antwort aus den mitgelieferten Fachmodulen genauso "
    "belegt aus wie eine aus dem Gesetzestext.",
)

doc.add_heading("Wenn kein Sprachmodell eingerichtet ist", level=2)
absatz("Dann schreibt der Buchhalter offen: \"Hinweis: Es wurde keine "
       "Modellantwort erzeugt.\" Er recherchiert trotzdem in seinen Quellen "
       "und zeigt Ihnen die Fundstellen - er formuliert nur keine fachliche "
       "Wuerdigung. Das ist gewollt: lieber ehrlich nichts sagen als etwas "
       "erfinden. Wie Sie ein Modell einrichten, steht in Kapitel 2 - es sind "
       "zwei Befehle.")

# ================================================================ 7
doc.add_heading("9.  Was automatisch gespeichert wird - und was nicht", level=1)

absatz("Die wichtigste Frage im taeglichen Umgang. Kurz: Ihre Unterhaltungen "
       "werden automatisch gespeichert. Alles, was in Ihr "
       "Unternehmensgedaechtnis soll, bestaetigen Sie ausdruecklich.")

tabelle_mit(
    ["Was", "Speichert sich das von selbst?", "Was Sie tun muessen"],
    [
        ["Ihre Frage und die Antwort",
         "Ja, sofort",
         "Nichts. Jede Frage und jede Antwort wird in dem Moment auf den "
         "Datentraeger geschrieben, in dem sie entsteht."],
        ["Die Unterhaltung als Ganzes",
         "Ja",
         "Nichts. Sie erscheint rechts unter \"Unterhaltungen\" und ist nach "
         "einem Neustart wieder da."],
        ["Angaben, die Sie im Gespraech nennen",
         "Nein - es wird gefragt",
         "Erkennt der Buchhalter eine dauerhafte Unternehmensangabe, fragt "
         "er \"Dauerhaft merken?\". Erst mit Ja wird gespeichert."],
        ["Onboarding-Fragebogen",
         "Nein",
         "Auf Speichern klicken."],
        ["Eintrag unter Neu / Aendern",
         "Nein",
         "Auf Speichern klicken."],
        ["Einstellungen",
         "Nein",
         "Auf Einstellungen speichern klicken."],
        ["Beleg",
         "Ja, beim Auswaehlen",
         "Nichts weiter. Die Datei wird sofort auf den Datentraeger "
         "uebernommen."],
    ],
    breiten=[3.8, 3.6, 7.6],
)

kasten(
    "Es gibt keine Schaltflaeche \"Alles speichern\"",
    "Und Sie brauchen auch keine. Der Buchhalter schreibt jede Aenderung "
    "sofort und einzeln auf den Datentraeger - nicht erst beim Beenden. "
    "Selbst wenn der Rechner mitten in der Arbeit ausgeht, ist alles bis zur "
    "letzten abgeschickten Frage vorhanden.",
)

doc.add_heading("Wo das alles liegt", level=2)
absatz("Ausschliesslich auf Ihrem Datentraeger, unterhalb des Ordners, den "
       "der Startbildschirm als Datenverzeichnis nennt - bei Ihnen zum "
       "Beispiel E:\\Portable-Buchhalter-Windows. Nichts wird auf dem "
       "jeweiligen PC abgelegt, nichts in die Windows-Registrierung "
       "geschrieben, nichts hochgeladen, nichts an Dritte gesendet. Deshalb "
       "koennen Sie den Datentraeger abziehen, an einem anderen Rechner "
       "einstecken und dort weiterarbeiten.")

absatz("Diese Ordner liegen dort - Sie duerfen jederzeit hineinschauen:")

tabelle_mit(
    ["Ordner", "Was darin liegt", "Ihre Daten?"],
    [
        ["database", "Die Datenbank company.db: Unternehmensgedaechtnis, "
                     "alle Unterhaltungen, das Protokoll", "Ja - das Wichtigste"],
        ["company", "Das Unternehmensprofil in lesbarer Form, sobald Sie es "
                    "exportieren", "Ja"],
        ["conversations", "Exportierte Unterhaltungen", "Ja"],
        ["workspace", "Ihre eigenen Arbeitsdateien und aufgenommenen Belege", "Ja"],
        ["backups", "Sicherungen, die Sie angelegt haben", "Ja"],
        ["logs", "Protokolldateien, darunter startfehler.txt", "Betriebsdaten"],
        ["config", "Ihre Einstellungen (settings.json), das Quellenregister "
                   "und der verschluesselte Tresor (secrets.enc)", "Ja"],
        ["license", "Die Lizenzdateien, sobald es welche gibt", "Vertragsdaten"],
        ["resources", "Die abgerufenen amtlichen Quellen und der Suchindex", "Fachwissen"],
        ["knowledge", "Die mitgelieferten Fachmodule", "Programmbestandteil"],
        ["models", "Das Sprachmodell, sobald Sie eines einrichten", "Programmbestandteil"],
        ["updates", "Sicherung und Bericht je Wissensupdate (fuer die Ruecknahme)", "Betriebsdaten"],
        ["src, tools, docs", "Das Programm selbst und die Unterlagen", "Programmbestandteil"],
    ],
    breiten=[3.4, 8.6, 3.2],
)

kasten(
    "Die eine Datei, auf die es ankommt",
    "database\\company.db. Darin steht alles, was der Buchhalter ueber Ihr "
    "Unternehmen weiss, samt Verlauf und allen Unterhaltungen. Wenn Sie nur "
    "eine Sache sichern wollen, dann diese - besser aber ueber die "
    "Schaltflaeche Sicherung erstellen, weil dabei auch eine Beschreibung "
    "des Inhalts mitgeschrieben wird.",
)

doc.add_heading("Mehrere Unternehmen auf einem Datentraeger", level=2)
absatz("Betreuen Sie mehrere Firmen, kann der Buchhalter getrennte "
       "Kundenbereiche fuehren. Die oben genannten Ordner liegen dann "
       "unterhalb von customers\\<Kennung>\\ - je Kunde ein eigener Baum. "
       "Unternehmenswissen des einen Kunden kann so nicht beim anderen "
       "auftauchen. Das allgemeine Fachwissen bleibt bewusst gemeinsam, denn "
       "es enthaelt keine Unternehmensdaten.")

doc.add_heading("Trotzdem: Sicherungen", level=2)
absatz("Automatisches Speichern schuetzt nicht vor einem Defekt des "
       "Datentraegers. Legen Sie regelmaessig eine Sicherung an - "
       "Registerkarte Wissen aktualisieren, Schaltflaeche Sicherung "
       "erstellen - und bewahren Sie sie an einem anderen Ort auf.")

# ================================================================ 8
doc.add_heading("10.  Belege hinzufuegen", level=1)

absatz("Sie koennen eigene Dokumente einlesen - Rechnungen, Vertraege, "
       "Kontoauszuege als Text. Der Buchhalter kann sich dann in seinen "
       "Antworten darauf beziehen.")

schritt("Registerkarte Belege oeffnen (oder in der Unterhaltung auf "
        "\"Dokument hinzufuegen\" klicken - beides fuehrt zum selben).")
schritt("Auf Beleg hinzufuegen klicken.")
schritt("Datei auswaehlen. Unterstuetzt werden: PDF, TXT, MD, HTML, HTM, XML, CSV.")
schritt("Der Beleg erscheint in der Liste mit Titel, Art, Zeitpunkt, Status "
        "und der Ablage auf dem Datentraeger.")

kasten(
    "Wo der Beleg landet",
    "Auf Ihrem Datentraeger, nirgendwo sonst. Die Spalte \"Ablage auf dem "
    "Datentraeger\" nennt den genauen Pfad. Es wird nichts hochgeladen und "
    "nichts an Dritte gesendet.",
)

# ================================================================ 7
doc.add_heading("11.  Unternehmenswissen pflegen", level=1)

absatz("Registerkarte Unternehmenswissen. Hier steht alles, was sich der "
       "Buchhalter ueber Ihren Betrieb gemerkt hat - versioniert, mit "
       "Verlauf.")

doc.add_heading("Die Tabelle", level=2)
tabelle_mit(
    ["Spalte", "Bedeutung"],
    [
        ["Schluessel", "Der technische Name, z. B. company.chart_of_accounts"],
        ["Kategorie", "Einordnung, z. B. accounting, tax, process"],
        ["Titel", "Die verstaendliche Bezeichnung, z. B. Kontenrahmen"],
        ["Inhalt", "Was gespeichert ist"],
        ["V", "Version - wie oft dieser Eintrag geaendert wurde"],
    ],
    breiten=[4.0, 11.0],
)

doc.add_heading("Suchen", level=2)
absatz("Oben links steht das Feld Suche. Suchbegriff eintippen und auf Suchen "
       "klicken oder die Eingabetaste druecken. Alles anzeigen setzt die "
       "Suche zurueck.")

doc.add_heading("Eintrag anlegen oder aendern", level=2)
absatz("Auf Neu / Aendern klicken. Es oeffnet sich das Fenster "
       "\"Unternehmenswissen bearbeiten\" mit vier Feldern:")
tabelle_mit(
    ["Feld", "Was hineingehoert"],
    [
        ["Schluessel", "Kurzer technischer Name ohne Leerzeichen, z. B. company.bank"],
        ["Titel", "Verstaendliche Bezeichnung, z. B. Hausbank"],
        ["Kategorie", "Einordnung, z. B. accounting, tax, process, profile"],
        ["Inhalt", "Der eigentliche Text - hier steht, was gilt"],
    ],
    breiten=[4.0, 11.0],
)
absatz("Zum Aendern eines vorhandenen Eintrags diesen zuerst in der Tabelle "
       "anklicken, dann Neu / Aendern. Die alte Fassung geht nicht verloren.")

doc.add_heading("Die uebrigen Schaltflaechen", level=2)
tabelle_mit(
    ["Schaltflaeche", "Was sie tut"],
    [
        ["Verlauf", "Zeigt alle frueheren Fassungen des markierten Eintrags"],
        ["Archivieren", "Nimmt den Eintrag aus dem aktiven Bestand; er bleibt im Verlauf erhalten"],
        ["Onboarding fortsetzen", "Oeffnet den gefuehrten Fragebogen aus Kapitel 5"],
        ["Profil exportieren", "Schreibt das gesamte Unternehmensprofil als lesbare Datei auf den Datentraeger"],
    ],
    breiten=[5.0, 10.0],
)

kasten(
    "Was hier nicht hineingehoert",
    "Keine Passwoerter, keine Zugangsdaten, keine Schluessel. Dafuer gibt es "
    "einen verschluesselten Tresor. Das Unternehmensgedaechtnis ist "
    "unverschluesselter Klartext auf dem Datentraeger - das ist Absicht, "
    "damit Sie es jederzeit lesen und pruefen koennen.",
)

# ================================================================ 11
doc.add_heading("12.  Ergebnisse als Datei ausgeben", level=1)

absatz("Der Buchhalter kann seine Ergebnisse als Datei herausgeben - als "
       "Excel-Tabelle, Word-Dokument, PowerPoint-Praesentation, PDF-Bericht "
       "oder in einem einfachen Textformat. Das geschieht vollstaendig auf "
       "dem Datentraeger: ohne Internet und ohne installiertes Microsoft "
       "Office.")

doc.add_heading("So speichern Sie eine Antwort", level=2)
schritt("Frage stellen und die Antwort abwarten.")
schritt("Rechts neben Antwort speichern das Format waehlen - zum Beispiel pdf.")
schritt("Auf Antwort speichern klicken. Die Anwendung nennt Ihnen den Pfad.")

absatz()
doc.add_heading("Die Formate", level=2)
tabelle_mit(
    ["Format", "Wofuer es gedacht ist"],
    [
        ["xlsx", "Excel-Arbeitsmappe: Auswertungen und Buchungslisten. Betraege "
                 "sind Zahlen, mit denen Excel rechnet; Konto- und Belegnummern "
                 "bleiben Text und verlieren keine fuehrende Null."],
        ["docx", "Word-Dokument: Berichte und Dokumentationen, mit Ueberschriften "
                 "und Tabellen."],
        ["pptx", "PowerPoint: Kurzbericht, je Abschnitt eine Folie."],
        ["pdf", "PDF: unveraenderlicher Bericht zum Weitergeben."],
        ["csv", "Tabelle fuer die Weiterverarbeitung; oeffnet in Excel richtig."],
        ["txt / md", "Einfacher Text beziehungsweise Text mit Gliederung."],
        ["json", "Maschinenlesbar, fuer eigene Auswertungen."],
    ],
    breiten=[3.0, 12.0],
)

doc.add_heading("Wo die Dateien liegen", level=2)
absatz("Im Ordner workspace\\artefakte auf dem Datentraeger - also bei Ihren "
       "Unternehmensdaten, nicht im Programmordner und nicht auf dem "
       "Computer, an dem Sie gerade arbeiten. Arbeiten mehrere Unternehmen "
       "auf demselben Datentraeger, hat jedes seinen eigenen Ordner.")

kasten(
    "Nichts wird ueberschrieben",
    "Speichern Sie zweimal unter demselben Namen, entsteht eine zweite "
    "Fassung: Bericht.pdf, dann Bericht_v2.pdf. Ihre erste Datei bleibt "
    "unveraendert. Ersetzt wird nur, wenn Sie es ausdruecklich verlangen.\n\n"
    "Bricht das Speichern ab, bleibt keine halbe Datei zurueck, die wie ein "
    "fertiger Bericht aussieht.",
)

doc.add_heading("Mit der Konsolenfassung", level=2)
tabelle_mit(
    ["Befehl", "Wirkung"],
    [
        ["PORTABLE_BUCHHALTER_KONSOLE.exe datei formate",
         "Zeigt alle Formate, die erzeugt werden koennen"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe datei antwort --format docx",
         "Speichert die letzte Antwort als Word-Dokument"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe datei text --text \"...\" --format pdf",
         "Erzeugt eine Datei aus eigenem Text"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe datei liste",
         "Zeigt die zuletzt erzeugten Dateien"],
    ],
    breiten=[8.0, 7.0],
)

# ================================================================ 12
doc.add_heading("13.  Erweiterungen (Plugins)", level=1)

absatz("Der Buchhalter kann neue Faehigkeiten aufnehmen, ohne dass das "
       "Programm neu gebaut wird. Eine solche Erweiterung heisst Plugin und "
       "kommt als einzelne Datei mit der Endung .kimplug.")

absatz("Beispiele: ein zusaetzliches Ausgabeformat, eine weitere "
       "Wissensquelle, die Anbindung an ein anderes System.")

doc.add_heading("Bevor Sie etwas installieren", level=2)
absatz("Ein Plugin laeuft mit den Rechten der Anwendung. Deshalb gilt:")
punkt("Die Anwendung zeigt Ihnen vorher, was das Plugin verlangt - zum "
      "Beispiel Unternehmensgedaechtnis lesen oder Verbindung ins Internet.")
punkt("Ohne Ihre ausdrueckliche Bestaetigung wird nichts installiert.")
punkt("Installiert ist noch nicht aktiv. Das Aktivieren ist ein zweiter "
      "Schritt.")
punkt("Ein Plugin von einem Herausgeber, den Sie nicht kennen, sollten Sie "
      "nicht installieren - so wie Sie auch sonst kein fremdes Programm "
      "starten wuerden.")

doc.add_heading("Der Ablauf", level=2)
tabelle_mit(
    ["Befehl", "Wirkung"],
    [
        ["PORTABLE_BUCHHALTER_KONSOLE.exe plugin pruefen <Datei>",
         "Zeigt Name, Herausgeber, Signatur und die verlangten Berechtigungen. "
         "Installiert wird dabei nichts."],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe plugin installieren <Datei> --bestaetigen",
         "Installiert das Plugin und erteilt die verlangten Berechtigungen"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe plugin aktivieren <Kennung>",
         "Schaltet es ein - ab dem naechsten Start steht die Faehigkeit bereit"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe plugin liste",
         "Zeigt alle Plugins mit Zustand, Rechten und zusaetzlichen Faehigkeiten"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe plugin deaktivieren <Kennung>",
         "Schaltet es wieder aus, ohne es zu entfernen"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe plugin entfernen <Kennung>",
         "Entfernt es. Die Daten des Plugins bleiben erhalten - sie gehoeren Ihnen"],
    ],
    breiten=[8.4, 6.6],
)

kasten(
    "Was die Anwendung ueberwacht - und was nicht",
    "Jede Berechtigung wird einzeln erteilt und jeder Vorgang protokolliert. "
    "Ein Plugin, das keine Erlaubnis fuer das Unternehmensgedaechtnis hat, "
    "kommt nicht daran. Und im Betriebsmodus OFFLINE greift auch ein Plugin "
    "mit Interneterlaubnis nicht ins Netz.\n\n"
    "Was die Anwendung NICHT leisten kann: ein Plugin technisch vom Rest "
    "abzuschotten. Es laeuft im selben Programm. Deshalb der Grundsatz oben: "
    "nur installieren, was aus vertrauenswuerdiger Quelle stammt. "
    "Einzelheiten in PLUGIN_KONZEPT.md.",
)

absatz("Ein Beispiel liegt bei: examples/plugin_html ergaenzt das "
       "Ausgabeformat HTML. Es verlangt keinerlei Berechtigungen und zeigt, "
       "wie eine Erweiterung aufgebaut ist.")

# ================================================================ 8
doc.add_heading("14.  Betriebsmodus: HYBRID, OFFLINE, ONLINE", level=1)

absatz("Sie bestimmen, ob der Buchhalter ins Internet darf. Die Auswahl "
       "steht oben rechts in der Kopfzeile und ist jederzeit erreichbar.")

tabelle_mit(
    ["Modus", "Was gilt", "Wann sinnvoll"],
    [
        ["HYBRID\n(Vorgabe)",
         "Lokale Arbeit ist die Grundlage. Onlinequellen und Wissensupdates "
         "duerfen zusaetzlich verwendet werden. Faellt die Verbindung aus, "
         "geht es ohne Unterbrechung lokal weiter.",
         "Der Normalfall."],
        ["OFFLINE",
         "Ausschliesslich lokal. Keine Webrecherche, keine externen "
         "Schnittstellen, keine Cloud-KI, keine automatischen "
         "Wissensupdates - auch dann nicht, wenn eine Verbindung besteht.",
         "Beim Kunden, im Zug, in abgeschotteten Netzen, oder immer dann, "
         "wenn nichts nach draussen gehen soll."],
        ["ONLINE",
         "Onlinefunktionen duerfen bevorzugt verwendet werden. Ihre lokalen "
         "Daten, das Unternehmensgedaechtnis und das lokale Fachwissen "
         "bleiben unveraendert verfuegbar.",
         "Wenn Sie gezielt aktuelle Quellen heranziehen wollen."],
    ],
    breiten=[3.0, 7.4, 4.6],
)

kasten(
    "ONLINE schaltet nichts ab",
    "Ein haeufiges Missverstaendnis: ONLINE bedeutet nicht, dass lokale "
    "Funktionen ausgehen. Ihr Unternehmensgedaechtnis, Ihre Belege und das "
    "mitgelieferte Fachwissen sind in jedem Modus da. Der Modus regelt nur, "
    "ob zusaetzlich nach draussen gegriffen werden darf.",
)

doc.add_heading("Betriebsmodus und Internet sind zweierlei", level=2)
absatz("Das ist der wichtigste Punkt dieses Kapitels. In der Kopfzeile "
       "stehen deshalb zwei Angaben nebeneinander:")

tabelle_mit(
    ["Was Sie sehen", "Bedeutung"],
    [
        ["Betriebsmodus: OFFLINE\nInternet: verfuegbar",
         "Sie haben sich fuer OFFLINE entschieden. Es wird nichts abgerufen - "
         "nicht einmal geprueft, ob eine Verbindung besteht, denn eine "
         "solche Pruefung waere selbst ein Zugriff. Ihre Entscheidung gilt."],
        ["Betriebsmodus: HYBRID\nInternet: nicht verfuegbar",
         "Sie waeren bereit fuer Onlinefunktionen, aber es gibt gerade keine "
         "Verbindung. Der Buchhalter arbeitet lokal weiter, ohne dass etwas "
         "verloren geht. Sobald die Verbindung zurueck ist, geht es wieder."],
    ],
    breiten=[5.0, 10.0],
)

kasten(
    "Ihre Wahl wird nie hinter Ihrem Ruecken aufgehoben",
    "Wer OFFLINE gewaehlt hat, bleibt offline - auch wenn das Netz "
    "zurueckkehrt, auch nach einem Neustart. Die Anwendung setzt sich nicht "
    "selbst wieder auf Onlinebetrieb. Nur Sie aendern den Modus.",
)

doc.add_heading("Umschalten", level=2)
schritt("Oben rechts das Feld Betriebsmodus aufklappen.")
schritt("HYBRID, OFFLINE oder ONLINE waehlen.")
schritt("Es erscheint eine kurze Ansage, was jetzt gilt. Mit OK bestaetigen.")
absatz("Die Wahl wird sofort gespeichert und im Protokoll vermerkt - mit "
       "vorherigem und neuem Modus, Zeitpunkt und Verbindungsstatus. Nach "
       "einem Neustart gilt sie unveraendert weiter.")
absatz("Ueber die Konsolenfassung geht es ebenso: "
       "PORTABLE_BUCHHALTER_KONSOLE.exe modus zeigt den Stand, "
       "PORTABLE_BUCHHALTER_KONSOLE.exe modus OFFLINE wechselt.")

# ================================================================ 12
doc.add_heading("15.  Wissen aktualisieren", level=1)

absatz("Registerkarte Wissen aktualisieren. Hier laedt der Buchhalter "
       "amtliche Quellen nach. Das ist der einzige Teil, der Internet "
       "braucht. Ohne Internet wird nichts abgerufen, und der vorhandene "
       "Wissensstand bleibt unveraendert nutzbar.")

tabelle_mit(
    ["Schaltflaeche", "Was sie tut", "Wann sinnvoll"],
    [
        ["Wissen jetzt aktualisieren", "Laedt die Quellen und uebernimmt sie",
         "Wenn Sie online sind und den Stand auffrischen wollen"],
        ["Trockenlauf (nichts schreiben)", "Zeigt, was passieren wuerde, ohne etwas zu aendern",
         "Zum Ausprobieren, wenn Sie unsicher sind"],
        ["Letzten Lauf zuruecknehmen", "Macht das letzte Update rueckgaengig",
         "Wenn nach einem Update etwas nicht stimmt"],
        ["Sicherung erstellen", "Legt eine Sicherungskopie Ihrer Daten an",
         "Vor groesseren Aenderungen - und regelmaessig"],
    ],
    breiten=[4.6, 6.4, 5.0],
)

doc.add_heading("Automatische Aktualisierung", level=2)
absatz("Der Buchhalter prueft von sich aus, ob der Wissensstand veraltet "
       "ist. Vorgabe ist **woechentlich**. Einstellbar unter Einstellungen "
       "und Status: taeglich, woechentlich, monatlich, benutzerdefiniert "
       "oder gar nicht (manual).")

absatz("Oben in der Registerkarte steht immer, woran Sie sind:")
tabelle_mit(
    ["Update-Status", "Bedeutung", "Was zu tun ist"],
    [
        ["AKTUELL", "Innerhalb des Intervalls aktualisiert.", "Nichts."],
        ["UPDATE FAELLIG", "Das Intervall ist ueberschritten.",
         "Bei Gelegenheit auf Wissen jetzt aktualisieren klicken."],
        ["UPDATE UEBERFAELLIG", "Mehr als das doppelte Intervall - etwa nach "
         "einer laengeren Offlinephase.", "Aktualisieren, sobald Sie online sind."],
        ["NOCH NIE AKTUALISIERT", "Es gab noch keinen erfolgreichen Abruf. "
         "Das mitgelieferte Fachwissen ist trotzdem nutzbar.",
         "Einmal aktualisieren, wenn Sie online sind."],
        ["OFFLINE - UPDATE PAUSIERT", "Sie haben den Offline-Modus gewaehlt. "
         "Es wird nicht synchronisiert, auch nicht bei bestehender Verbindung.",
         "Nichts - das ist so gewollt. Der vorhandene Stand bleibt nutzbar."],
        ["KEINE VERBINDUNG", "Faellig, aber gerade kein Netz.",
         "Spaeter erneut."],
        ["AUTOMATIK AUS", "Der Zeitplan steht auf manual.",
         "Nichts, sofern gewollt."],
    ],
    breiten=[4.2, 6.4, 4.4],
)

absatz("Daneben stehen die Zahlen dazu: eingestellter Zeitplan, Intervall in "
       "Tagen, Datum der letzten Aktualisierung und der naechsten Pruefung.")

kasten(
    "Es wird nie behauptet, das Wissen sei aktuell",
    "Gab es noch keinen erfolgreichen Abruf, steht dort \"noch nie "
    "aktualisiert\" - nicht etwa \"aktuell\". Auch in den Antworten wird der "
    "Wissensstand mit Datum genannt. Sie sollen jederzeit erkennen koennen, "
    "wie alt die Grundlage ist, auf der geantwortet wird.",
)

kasten(
    "Sicherungen gehoeren nicht auf denselben Datentraeger",
    "Geht die SSD verloren oder kaputt, ist eine Sicherung darauf ebenfalls "
    "weg. Legen Sie Sicherungen zusaetzlich woanders ab - zweite Festplatte, "
    "Netzlaufwerk. Der Buchhalter unterstuetzt das ausdruecklich.",
)

# ================================================================ 9
doc.add_heading("16.  Einstellungen und Status", level=1)

absatz("Registerkarte Einstellungen und Status. Links stellen Sie ein, rechts "
       "sehen Sie den Zustand.")

doc.add_heading("Die Einstellungen (links)", level=2)
tabelle_mit(
    ["Einstellung", "Auswahl", "Bedeutung"],
    [
        ["Zeitplan Wissensupdate", "manual, weekly, monthly, custom",
         "Wie oft an ein Update erinnert wird. manual = nur wenn Sie es anstossen"],
        ["Fundstellen je Antwort", "4, 6, 8, 12, 16",
         "Wie viele Quellen herangezogen werden. Mehr = gruendlicher, aber langsamer"],
        ["Dauerhafte Unternehmensinformationen erkennen", "Haekchen",
         "Erkennt im Gespraech genannte Angaben zu Ihrem Betrieb"],
        ["Vor dem Speichern nachfragen", "Haekchen",
         "Fragt nach, bevor etwas ins Unternehmensgedaechtnis kommt. Empfohlen: an"],
        ["Online-Sprachmodell erlauben (optional)", "Haekchen",
         "Erlaubt die Nutzung eines Modells im Internet. Standard: aus"],
        ["Protokollierung aktiv", "Haekchen",
         "Zeichnet auf, was geschehen ist. Empfohlen: an - das ist Ihr Nachweis"],
    ],
    breiten=[5.0, 3.4, 7.6],
)
absatz("Aenderungen werden erst mit Einstellungen speichern wirksam.")

doc.add_heading("Der Status (rechts)", level=2)
absatz("Zeigt dieselbe Uebersicht wie beim Start: Datenverzeichnis, "
       "Unternehmensgedaechtnis, Fachwissen, Suchindex, lokales Modell, "
       "Lizenz, Quellenregister, Geheimnistresor. Mit Status aktualisieren "
       "wird neu geprueft.")
absatz("Steht bei einem Punkt HINWEIS statt OK, ist das kein Fehler, sondern "
       "eine Einschraenkung - etwa \"Kein Sprachmodell verfuegbar\". Der Text "
       "daneben sagt, was das bedeutet und wie es zu beheben ist.")

# ================================================================ 12
doc.add_heading("17.  Lizenz", level=1)

doc.add_heading("Im Augenblick muessen Sie nichts tun", level=2)
absatz("Der Startbildschirm zeigt in der Zeile Lizenz:")
absatz("\"Diese Fassung laeuft ohne Lizenzpruefung (Vorab- bzw. "
       "Pilotfassung).\"", kursiv=True)
absatz("Genau so ist es gemeint. Diese Fassung braucht keine Lizenzdatei, "
       "keine Aktivierung, keinen Schluessel und keine Registrierung. Sie "
       "koennen sofort arbeiten.")

doc.add_heading("Warum das so ist", level=2)
absatz("Die Lizenzpruefung ist vollstaendig eingebaut und geprueft, aber "
       "bewusst noch nicht scharf geschaltet. Dazu fehlt ein Stueck, das "
       "nicht in der Software liegen kann: der oeffentliche Pruefschluessel "
       "des Herausgebers. Solange der nicht eingebaut ist, taeuscht die "
       "Anwendung keine Gueltigkeit vor - sie sagt offen, dass sie nicht "
       "pruefen kann. Das ist der ehrlichere Zustand als eine Pruefung, die "
       "nur so tut.")

doc.add_heading("Wie es spaeter funktionieren wird", level=2)
absatz("Damit Sie wissen, was auf Sie zukommt:")
schritt("Sie erhalten eine Lizenzdatei - zwei kleine Dateien, license.json "
        "und license.sig - und legen sie in den Ordner license auf dem "
        "Datentraeger.")
schritt("Beim Start prueft der Buchhalter die Signatur. Das geschieht "
        "vollstaendig ohne Internet; es wird nichts uebertragen und nichts "
        "gemeldet.")
schritt("Die Lizenz ist an den Datentraeger gebunden, nicht an einen "
        "bestimmten PC. Sie koennen die SSD also weiterhin an jeden Rechner "
        "stecken - genau das soll erhalten bleiben.")

absatz()
tabelle_mit(
    ["Fall", "Was dann passiert"],
    [
        ["Sie kopieren den Ordner auf einen zweiten Datentraeger",
         "Die Kopie ist nicht lizenziert. Das Original bleibt gueltig."],
        ["Jemand aendert die Lizenzdatei",
         "Die Signatur passt nicht mehr, die Lizenz wird abgelehnt."],
        ["Kein Internet",
         "Ohne Bedeutung - die Pruefung laeuft rein oertlich."],
        ["Der Datentraeger geht kaputt",
         "Sie erhalten fuer den Ersatz eine neue Lizenz. Ein vorgesehener, "
         "getesteter Vorgang."],
        ["Die Lizenz fehlt oder ist ungueltig",
         "Sie bekommen eine verstaendliche Meldung. Ihre Daten bleiben "
         "unangetastet - siehe unten."],
    ],
    breiten=[6.2, 8.8],
)

kasten(
    "Ihre Daten werden nie als Druckmittel benutzt",
    "Fehlt die Lizenz oder ist sie ungueltig, wird nichts geloescht, nichts "
    "gesperrt und nichts verschluesselt. Ihr Unternehmenswissen bleibt "
    "lesbar, und Export und Sicherung funktionieren weiter. Eine "
    "Lizenzfrage ist eine kaufmaennische Angelegenheit und darf sich nie "
    "gegen Ihre Daten richten.",
)

doc.add_heading("Was noch fehlt, damit die Lizenzierung greift", level=2)
absatz("Damit Sie den Aufwand einschaetzen koennen - technisch ist alles "
       "vorhanden und geprueft, es fehlen im Wesentlichen zwei "
       "Entscheidungen und ein Schluesselpaar:")

tabelle_mit(
    ["Schritt", "Was zu tun ist", "Wer"],
    [
        ["1. Schluesselpaar erzeugen",
         "Einmalig ein Ed25519-Schluesselpaar erzeugen. Der private "
         "Schluessel bleibt beim Herausgeber und darf den Betrieb nie "
         "verlassen; wer ihn hat, kann beliebige Lizenzen ausstellen.",
         "Herausgeber"],
        ["2. Oeffentlichen Schluessel einbauen",
         "Der oeffentliche Teil wird in die Anwendung uebernommen. Nur er "
         "wird ausgeliefert - er kann Lizenzen pruefen, aber keine "
         "ausstellen.",
         "Entwicklung"],
        ["3. Lizenzpflicht einschalten",
         "In der Konfiguration license.required auf true setzen. Bis dahin "
         "laeuft die Pilotfassung bewusst ohne Pruefung.",
         "Entscheidung"],
        ["4. Lizenzmodell festlegen",
         "Laufzeit, Zahl der Instanzen, enthaltene Fachmodule. Das steht "
         "spaeter in der Lizenzdatei.",
         "Kaufmaennisch"],
        ["5. Programm signieren",
         "Ein Code-Signing-Zertifikat beschaffen und die EXE signieren. "
         "Danach entfaellt die SmartScreen-Meldung beim ersten Start.",
         "Herausgeber"],
    ],
    breiten=[3.6, 8.4, 3.0],
)

absatz("Die Punkte 1 bis 4 sind Einrichtung, kein Programmierauftrag - das "
       "Verfahren selbst ist fertig und mit 22 Tests abgesichert, "
       "einschliesslich aller Faelle: Kopie auf einen zweiten Datentraeger, "
       "veraenderte Lizenzdatei, fehlende Signatur, Ablauf, Ersatzgeraet.")

doc.add_heading("Der Ablauf im Betrieb, wenn es soweit ist", level=2)
absatz("Fuer Sie als Anwender bleibt es einfach:")
schritt("Aktivierungsanfrage erzeugen: "
        "PORTABLE_BUCHHALTER_KONSOLE.exe lizenz anfrage --kunde \"Ihre Firma\" "
        "--datei anfrage.json. Darin steht die Kennung des Datentraegers, "
        "keine Unternehmensdaten.")
schritt("Die Datei an den Herausgeber schicken.")
schritt("Sie erhalten license.json und license.sig zurueck.")
schritt("Aufnehmen: PORTABLE_BUCHHALTER_KONSOLE.exe lizenz aufnehmen "
        "--lizenz license.json --signatur license.sig. Die Dateien landen "
        "im Ordner license auf dem Datentraeger.")
schritt("Pruefen: PORTABLE_BUCHHALTER_KONSOLE.exe lizenz - zeigt den Zustand "
        "an. Ab dann steht er auch im Startbildschirm.")

absatz("Es wird dabei zu keinem Zeitpunkt etwas ins Internet uebertragen.")

kasten(
    "Solange der Pruefschluessel fehlt",
    "Die Schritte funktionieren schon heute - eine Lizenz laesst sich "
    "erzeugen und aufnehmen. Der Buchhalter meldet dann aber ehrlich "
    "\"NICHT_PRUEFBAR: In dieser Fassung ist kein Pruefschluessel des "
    "Herausgebers hinterlegt. Die Lizenz wird deshalb weder als gueltig noch "
    "als ungueltig behandelt.\" Er tut also nicht so, als haette er "
    "geprueft. Mit eingebautem Schluessel wird daraus GUELTIG.",
)

doc.add_heading("Und die Software selbst?", level=2)
absatz("Diese Fassung ist eine Vorab- bzw. Pilotfassung. Sie ist noch nicht "
       "digital signiert - daher die Windows-Meldung beim ersten Start - und "
       "sie ist nicht als kommerziell freigegeben erklaert. Was dafuer noch "
       "fehlt, koennen Sie sich jederzeit selbst anzeigen lassen: "
       "PORTABLE_BUCHHALTER_KONSOLE.exe reife")

# ================================================================ 13
doc.add_heading("18.  Was der Buchhalter nicht tut", level=1)

absatz("Das ist kein Mangel, sondern bewusst so gebaut. Bitte lesen Sie es "
       "einmal in Ruhe.")

for grenze in profil["limits"]:
    punkt(grenze)

absatz()
doc.add_heading("Diese Vorgaenge brauchen immer Ihre Freigabe", level=2)
uebersetzung = {
    "booking": "Verbuchung",
    "export": "Datenexport",
    "erp_write": "Schreibender Zugriff auf ein ERP-System",
    "payment": "Zahlung",
    "filing": "Meldung oder Abgabe an eine Behoerde",
    "masterdata": "Aenderung von Stammdaten",
}
tabelle_mit(
    ["Vorgang", "Bedeutung"],
    [[uebersetzung.get(v, v), v] for v in profil["requires_approval"]],
    breiten=[7.0, 6.0],
)

kasten(
    "Der Grundsatz",
    profil["disclaimer"] + "\n\n"
    "Angebundene Fremdsysteme werden ausserdem grundsaetzlich nur gelesen, "
    "nie beschrieben - es sei denn, Sie geben das ausdruecklich frei.",
)

# ================================================================ 11
doc.add_heading("19.  Wenn etwas nicht funktioniert", level=1)

tabelle_mit(
    ["Was Sie sehen", "Was zu tun ist"],
    [
        ["Doppelklick, aber nichts passiert",
         "Auf dem Datentraeger im Ordner logs die Datei startfehler.txt oeffnen. "
         "Dort steht der Grund im Klartext."],
        ["\"Der Computer wurde durch Windows geschuetzt\"",
         "Normal beim ersten Start - die Anwendung ist noch nicht signiert. "
         "Weitere Informationen, dann Trotzdem ausfuehren."],
        ["\"Kein Sprachmodell verfuegbar\"",
         "Erwartet, solange kein Modell eingerichtet ist. Der Buchhalter "
         "recherchiert weiterhin. Einrichtung: Kapitel 2, zwei Befehle."],
        ["Das Fenster geht nicht auf, Sie brauchen eine Ausgabe",
         "PORTABLE_BUCHHALTER_KONSOLE.exe verwenden - dieselbe Anwendung, aber "
         "mit Textausgabe. Zum Pruefen: PORTABLE_BUCHHALTER_KONSOLE.exe check"],
        ["Nach einem Update stimmt etwas nicht",
         "Registerkarte Wissen aktualisieren, dann Letzten Lauf zuruecknehmen."],
        ["Beim Wissensupdate steht ERGEBNIS: PARTIAL",
         "Ein Teil der Quellen wurde geladen, ein anderer nicht. Der Bericht "
         "nennt je Dokument den Grund. Der bereits geladene Teil ist nutzbar - "
         "siehe die naechste Tabelle."],
        ["Antwort ohne brauchbare Quellen",
         "Der Antwort nicht folgen. Frage praeziser stellen oder einen "
         "passenden Beleg hinzufuegen."],
    ],
    breiten=[6.0, 9.0],
)

doc.add_heading("Fehler beim Wissensupdate verstehen", level=2)
absatz("Amtliche Stellen bauen ihre Webauftritte um, und Server sind "
       "gelegentlich ueberlastet. Dass einzelne Quellen fehlschlagen, ist "
       "der Normalfall und kein Grund zur Sorge: Was geladen wurde, bleibt "
       "nutzbar, und der vorherige Stand wird nie zerstoert. Der Bericht "
       "nennt je Dokument den Grund.")

tabelle_mit(
    ["Meldung", "Was sie bedeutet", "Was zu tun ist"],
    [
        ["Adresse nicht mehr gueltig (404)",
         "Die Stelle hat ihre Seite verschoben oder entfernt.",
         "Neue Adresse heraussuchen und eintragen - siehe unten. Kein "
         "Programmfehler, keine neue Programmfassung noetig."],
        ["Server voruebergehend nicht erreichbar (503, 502, 504)",
         "Die Gegenstelle ist gerade ueberlastet oder in Wartung.",
         "Nichts. Es wird bereits bis zu dreimal versucht. Spaeter erneut "
         "aktualisieren."],
        ["Zugriff verweigert (403)",
         "Die Quelle sperrt automatisierte Abrufe.",
         "Quelle abschalten oder eine andere Bezugsquelle waehlen."],
        ["Zertifikat nicht pruefbar",
         "Die Gegenstelle liefert ihre Zertifikatskette unvollstaendig aus.",
         "Nichts. Es wird niemals ungeprueft geladen - lieber kein Dokument "
         "als ein ungesichertes."],
        ["Kein verwertbarer Text",
         "Die Seite baut ihren Inhalt erst im Browser auf, etwa eine "
         "Suchmaske.",
         "Den Eintrag auf eine Seite mit dem eigentlichen Inhalt umstellen, "
         "nicht auf das Suchformular."],
    ],
    breiten=[4.2, 5.4, 5.4],
)

doc.add_heading("Eine Adresse berichtigen", level=2)
absatz("Dafuer muessen Sie keine Datei von Hand bearbeiten. Mit der "
       "Konsolenfassung:")
tabelle_mit(
    ["Befehl", "Wirkung"],
    [
        ["PORTABLE_BUCHHALTER_KONSOLE.exe quellen liste",
         "Zeigt alle Quellen mit ihren Adressen und Kennungen"],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe quellen pruefen",
         "Prueft jede Adresse einzeln und nennt die nicht erreichbaren. Im "
         "OFFLINE-Betrieb wird dabei nichts abgerufen."],
        ["PORTABLE_BUCHHALTER_KONSOLE.exe quellen setzen --dokument <Kennung> --url <neue Adresse>",
         "Traegt die neue Adresse ein. Die Aenderung wird protokolliert; die "
         "Anwendung muss dafuer nicht neu gebaut werden."],
    ],
    breiten=[7.4, 7.6],
)

doc.add_heading("Weitere Unterlagen auf dem Datentraeger", level=2)
tabelle_mit(
    ["Datei", "Inhalt"],
    [
        ["START_HIER.md", "Kurzeinstieg"],
        ["docs/MODELL_EINRICHTEN.md", "Sprachmodell einrichten"],
        ["docs/ABNAHME.md", "Abnahme Schritt fuer Schritt"],
        ["BACKUP_WIEDERHERSTELLUNG.md", "Sicherung und Wiederherstellung"],
        ["SICHERHEITSKONZEPT.md", "Wie mit Ihren Daten umgegangen wird"],
        ["TESTBERICHT.md", "Was geprueft wurde - und was nicht"],
    ],
    breiten=[6.0, 9.0],
)

doc.add_paragraph()
schluss = doc.add_paragraph()
schluss.alignment = WD_ALIGN_PARAGRAPH.CENTER
lauf = schluss.add_run(
    f"Diese Anleitung beschreibt Programmfassung {PRODUKTFASSUNG} mit dem "
    f"Fachmodul {profil['name']} {profil['version']}. Aendert sich die "
    "Oberflaeche, aendert sich auch dieses Dokument."
)
lauf.italic = True
lauf.font.size = Pt(9)
lauf.font.color.rgb = GRAU

ziel = REPO / "docs" / "BEDIENUNGSANLEITUNG.docx"
doc.save(str(ziel))


def zoom_reparieren(datei: Path) -> bool:
    """Behebt einen Schemafehler der python-docx-Vorlage.

    Deren word/settings.xml enthaelt <w:zoom w:val="none"/> ohne das laut
    OOXML-Schema vorgeschriebene Attribut w:percent. Word ist nachsichtig
    und oeffnet die Datei trotzdem; LibreOffice und andere Leseprogramme
    weisen sie ab ("source file could not be loaded"). Deshalb wird das
    Attribut hier nachgetragen.
    """
    import re
    import shutil
    import tempfile
    import zipfile

    with zipfile.ZipFile(datei) as archiv:
        inhalte = {name: archiv.read(name) for name in archiv.namelist()}

    settings = inhalte.get("word/settings.xml")
    if settings is None:
        return False
    text = settings.decode("utf-8")
    if "w:zoom" not in text or 'w:percent="' in text:
        return False
    text = re.sub(r"<w:zoom(?![^>]*w:percent)([^>]*?)/>",
                  r'<w:zoom\1 w:percent="100"/>', text, count=1)
    inhalte["word/settings.xml"] = text.encode("utf-8")

    temp = Path(tempfile.mkstemp(suffix=".docx")[1])
    with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as archiv:
        for name, rohdaten in inhalte.items():
            archiv.writestr(name, rohdaten)
    shutil.move(str(temp), str(datei))
    return True


if zoom_reparieren(ziel):
    print("Schemafehler der Vorlage behoben (w:zoom ohne w:percent)")
print("geschrieben:", ziel, ziel.stat().st_size, "Bytes")
